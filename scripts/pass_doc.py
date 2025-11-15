#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOC/DOCX pass:
 - Extracts native text from .docx (python-docx) and .doc (antiword -> catdoc fallback)
 - Computes reliability and writes a single CSV row (per-document)
 - Accepts if reliability >= cutoff; otherwise tries DOC->PDF->TXT fallback;
   only if that also fails does it return non-zero so the orchestrator can quarantine.

Args:
  1: path to DOC/DOCX file
  2: path to CSV for this run
  3: path to run.log

Env:
  PASS_DOC_CUTOFF (float, default 0.75)
  PASS_DOCX_CUTOFF (float, default 0.70)
  LOG_LEVEL (INFO/DEBUG), optional
"""
import os
import sys
import subprocess
from typing import Optional

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

import common  # shared logger, CsvWriter, score_reliability
import output_writer

# Optional imports for PDF fallback
try:
    import doc_to_pdf
except Exception:  # pragma: no cover
    doc_to_pdf = None

try:
    import pass_pdf_txt
except Exception:  # pragma: no cover
    pass_pdf_txt = None


def _env_float(name: str, default: float) -> float:
    try:
        return float(os.getenv(name, str(default)))
    except Exception:
        return default


def _docx_text(path: str) -> str:
    # Extract paragraphs + table cell text
    try:
        import docx  # python-docx
    except Exception as e:
        raise RuntimeError(f"python-docx not available: {e}")

    try:
        d = docx.Document(path)
    except Exception as e:
        raise RuntimeError(f"docx open failed: {e}")

    parts = []
    # paragraphs
    for p in d.paragraphs:
        if p.text:
            parts.append(p.text)
    # tables
    for t in d.tables:
        try:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        except Exception:
            # table iteration is best-effort
            pass
    return "\n".join(parts)


def _run_cmd(cmd: list[str]) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)


def _doc_text(path: str) -> str:
    # Prefer antiword; fallback to catdoc if present
    try:
        cp = _run_cmd(["antiword", path])
        if cp.returncode == 0 and cp.stdout:
            return cp.stdout
    except FileNotFoundError:
        pass  # antiword not installed

    # Fallback: catdoc
    try:
        cp = _run_cmd(["catdoc", path])
        if cp.returncode == 0 and cp.stdout:
            return cp.stdout
    except FileNotFoundError:
        pass

    raise RuntimeError("Neither antiword nor catdoc produced text")


def _fallback_via_pdf(
    in_path: str,
    csv_path: str,
    run_log_path: str,
    logger,
) -> bool:
    """
    Second-chance path for DOC/DOCX:
      1. Convert DOC/DOCX -> PDF (using doc_to_pdf).
      2. Run PDF TXT pass (per-doc) via pass_pdf_txt.run.
      3. If we get usable text, score it and write a TXT+CSV row using the original DOC/DOCX path.
    Returns True on success, False if fallback failed or unavailable.
    """
    if doc_to_pdf is None or pass_pdf_txt is None:
        logger.warning("PDF fallback unavailable: doc_to_pdf or pass_pdf_txt not importable.")
        return False

    try:
        pdf_path = doc_to_pdf.convert_to_pdf(in_path, logger=logger)
    except Exception as e:
        logger.error(f"DOC->PDF conversion error: {e}")
        return False

    if not pdf_path or not os.path.exists(pdf_path):
        logger.error("DOC->PDF conversion failed to produce a valid PDF file.")
        return False

    logger.info(f"PDF fallback: running TXT pass on {os.path.basename(pdf_path)}")

    try:
        txt_ok, txt_payload = pass_pdf_txt.run(
            pdf_path,
            mode="per-doc",
            cutoff=0.0,   # let reliability be handled here, we just want any text
            logger=logger,
        )
    except Exception as e:
        logger.error(f"PDF TXT fallback error: {e}")
        # best effort cleanup
        try:
            os.remove(pdf_path)
        except Exception:
            pass
        return False

    text = ""
    if txt_ok and txt_payload:
        # per-doc mode: expect "text" key
        text = txt_payload.get("text") or ""

    text = text or ""
    if not text.strip():
        logger.warning("PDF TXT fallback produced no usable text.")
        try:
            os.remove(pdf_path)
        except Exception:
            pass
        return False

    # Score the fallback text and write as a single-page doc
    rel2 = common.score_reliability(text)
    logger.info(f"PDF TXT fallback success: reliability={rel2:.2f}")

    pages = [(1, text)]
    output_writer.write_result(
        csv_path=csv_path,
        original_file=os.path.abspath(in_path),
        pages=pages,
        pass_used="doc_pdf_text",
        score=rel2,
        status="OK",
        used_ocr=False,
        logger=logger,
    )

    # Cleanup temporary PDF
    try:
        os.remove(pdf_path)
    except Exception:
        pass

    return True


def main():
    if len(sys.argv) < 4:
        print("usage: pass_doc.py <doc|docx_path> <csv_path> <run_log_path>", file=sys.stderr)
        sys.exit(2)

    in_path, csv_path, run_log_path = sys.argv[1], sys.argv[2], sys.argv[3]
    basename = os.path.basename(in_path)
    ext = os.path.splitext(in_path)[1].lower()

    logger = common.get_logger(run_log_path)

    # Separate, tunable cutoffs for DOC and DOCX
    base_cutoff = _env_float("PASS_DOC_CUTOFF", 0.75)
    docx_cutoff = _env_float("PASS_DOCX_CUTOFF", 0.70)  # slightly more lenient by default

    # Extract natively (DOC/DOCX)
    try:
        if ext == ".docx":
            method = "docx_text"
            text = _docx_text(in_path)
            cutoff = docx_cutoff
        elif ext == ".doc":
            method = "doc_text"
            text = _doc_text(in_path)
            cutoff = base_cutoff
        else:
            logger.error(f"pass_doc called with unsupported extension: {ext}")
            sys.exit(2)
    except Exception as e:
        logger.error(f"DOC open/extract failed: {basename} :: {e}")
        # Record failure in CSV (no txt file)
        output_writer.write_result(
            csv_path=csv_path,
            original_file=os.path.abspath(in_path),
            pages=[],
            pass_used="doc_extract_error",
            score=0.0,
            status="ERROR",
            used_ocr=False,
            logger=logger,
        )
        sys.exit(1)

    text = text or ""
    rel = common.score_reliability(text)
    logger.info(f"DOC summary: {basename} method={method} reliability={rel:.2f} cutoff={cutoff}")

    # Gate: require some non-whitespace text and reliability above cutoff
    if text.strip() and rel >= cutoff:
        pages = [(1, text)]
        logger.info(f"DOC accept (native): {basename} reliability={rel:.2f}")

        output_writer.write_result(
            csv_path=csv_path,
            original_file=os.path.abspath(in_path),
            pages=pages,
            pass_used=method,
            score=rel,
            status="OK",
            used_ocr=False,
            logger=logger,
        )
        sys.exit(0)

    # Below cutoff or effectively empty: try PDF fallback first
    logger.warning(
        f"DOC/DOCX below cutoff or empty: {basename} reliability={rel:.2f} < {cutoff}. "
        "Attempting DOC->PDF TXT fallback."
    )

    if _fallback_via_pdf(in_path, csv_path, run_log_path, logger):
        logger.info(f"DOC/DOCX PDF fallback accepted: {basename}")
        sys.exit(0)

    # Fallback failed: mark as ERROR, no txt file
    logger.warning(f"DOC/DOCX fallback failed: {basename} -- recording ERROR")
    output_writer.write_result(
        csv_path=csv_path,
        original_file=os.path.abspath(in_path),
        pages=[],  # no usable text => no .txt
        pass_used=method,
        score=rel,
        status="ERROR",
        used_ocr=False,
        logger=logger,
    )
    sys.exit(1)


if __name__ == "__main__":
    main()
